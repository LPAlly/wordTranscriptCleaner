from docx import Document
import numpy as np
import argparse

def getText(filename):
	doc = Document(filename)
	fullText = []

	for para in doc.paragraphs:
		fullText.append(para.text)

	fullTextnp = np.array(fullText)
	indexOfText = fullTextnp != ""

	return(fullTextnp[indexOfText])

def getTalkingPersonParagraph(para):
	persons = []
	lastTalkingPerson = ""
	personTalkingOrder = []

	for index, paragraph in enumerate(para):
		delimPerson = paragraph.find(":")

		if delimPerson != -1:
			person = paragraph[0:delimPerson].strip()

			if person not in persons:
				persons.append(person)

			lastTalkingPerson = person

			personTalkingOrder.append(lastTalkingPerson)
		else:
			personTalkingOrder.append(lastTalkingPerson)

	return personTalkingOrder

def concatParagraph(para, talkingOrder):
	talkingPerson = talkingOrder[0]
	newPara = talkingPerson + " : "

	newDocument = Document()	

	for indexPara, person in enumerate(talkingOrder):
		if talkingPerson != person:
			newDocument.add_paragraph(newPara)
			newPara = ""
			newPara = person + " :"
			talkingPerson = person

		if para[indexPara].find(person) == -1:
			newPara += " " + para[indexPara]
		else:
			checkIfText = para[indexPara].split(":",1)
			if len(checkIfText) > 1:
				newPara += " " + checkIfText[1].strip()

	newDocument.add_paragraph(newPara)
	newDocument.save("documentCorrige.docx")

if __name__ == '__main__':
	parser = argparse.ArgumentParser()
	parser.add_argument('--filename', type=str, required=True)
	args = parser.parse_args()

	docxPath = args.filename

	paragraphs = getText(docxPath)
	talkingOrder = getTalkingPersonParagraph(paragraphs)
	concatParagraph(paragraphs, talkingOrder)
